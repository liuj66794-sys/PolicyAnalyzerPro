from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from core.config import AppConfig, DEFAULT_CONFIG, apply_tesseract_runtime_environment
from core.import_preview import ImportPreviewState
from core.text_cleaner import TextCleaner


class DocumentImportError(RuntimeError):
    pass


@dataclass(slots=True)
class PdfImportOptions:
    ocr_page_spec: str = ""
    use_ocr_cache: bool = True


@dataclass(slots=True)
class OcrLoadResult:
    text: str
    page_numbers: list[int]
    page_range_label: str
    cache_hit: bool = False


class DocumentLoader:
    TEXT_SUFFIXES = {".txt", ".md", ".log"}
    DOCX_SUFFIXES = {".docx"}
    PDF_SUFFIXES = {".pdf"}
    SUPPORTED_SUFFIXES = TEXT_SUFFIXES | DOCX_SUFFIXES | PDF_SUFFIXES
    _CJK_TOKEN = r"\u3400-\u9fff\uf900-\ufaff"
    _NUMERAL_TOKEN = r"0-9〇零一二三四五六七八九十百千万亿两"
    _PDF_NOISE_PATTERNS = (
        re.compile(r"^\d{4}[-/]\d{1,2}[-/]\d{1,2}\s+\d{1,2}:\d{2}(?::\d{2})?\b.*$"),
        re.compile(r"^https?://\S+.*$", re.IGNORECASE),
        re.compile(r"^www\.\S+$", re.IGNORECASE),
        re.compile(r"^\d+\s*/\s*\d+\s*$"),
    )
    _PDF_HEADING_PATTERNS = (
        re.compile(r"^[一二三四五六七八九十]+[、.]"),
        re.compile(r"^（[一二三四五六七八九十]+）"),
        re.compile(r"^各位代表[:：]"),
    )
    _AUTHOR_ROLE_MARKERS = (
        "国务院总理",
        "国务院副总理",
        "全国人大常委会委员长",
        "全国政协主席",
        "全国政协副主席",
        "国家主席",
        "最高人民法院院长",
        "最高人民检察院检察长",
        "部长",
        "主任",
        "秘书长",
    )

    def __init__(self, config: AppConfig | None = None) -> None:
        self.config = config or DEFAULT_CONFIG
        self._preview_cleaner = TextCleaner(self.config)
        self.last_load_state = ImportPreviewState()

    def reset_last_load_state(self) -> None:
        self.last_load_state = ImportPreviewState()

    def _update_last_load_state(
        self,
        path: Path,
        content: str,
        extraction_mode: str,
        *,
        ocr_page_range: str = "",
        ocr_page_count: int = 0,
        ocr_cache_hit: bool = False,
    ) -> None:
        normalized = content.replace("\r\n", "\n").replace("\r", "\n")
        non_empty_line_count = len([line for line in normalized.splitlines() if line.strip()])
        cleaned_paragraph_count = len(self._preview_cleaner.clean_paragraphs(content))
        abnormal_blank_lines = bool(re.search(r"\n\s*\n\s*\n+", normalized))
        self.last_load_state = ImportPreviewState(
            source_path=str(path),
            source_suffix=path.suffix.lower(),
            extraction_mode=extraction_mode,
            raw_char_count=len(content.strip()),
            non_empty_line_count=non_empty_line_count,
            cleaned_paragraph_count=cleaned_paragraph_count,
            abnormal_blank_lines=abnormal_blank_lines,
            ocr_page_range=ocr_page_range,
            ocr_page_count=ocr_page_count,
            ocr_cache_hit=ocr_cache_hit,
        )

    @classmethod
    def file_dialog_filter(cls) -> str:
        return (
            "文档文件 (*.txt *.md *.log *.docx *.pdf);;"
            "文本文件 (*.txt *.md *.log);;"
            "Word 文档 (*.docx);;"
            "PDF 文件 (*.pdf);;"
            "所有文件 (*.*)"
        )

    def load_text_from_path(
        self,
        path: str | Path,
        pdf_options: PdfImportOptions | None = None,
    ) -> str:
        self.reset_last_load_state()
        
        # 增强输入验证
        if path is None or (isinstance(path, str) and not path.strip()):
            raise DocumentImportError("\u6587\u6863\u8def\u5f84\u4e0d\u80fd\u4e3a\u7a7a")
        
        file_path = Path(path)
        
        # 验证文件路径安全性
        try:
            # 解析路径，防止路径遍历攻击
            file_path = file_path.resolve(strict=True)
        except Exception as exc:
            raise DocumentImportError(f"\u6587\u6863\u8def\u5f84\u89e3\u6790\u5931\u8d25\uff1a{exc}")
        
        if not file_path.exists():
            raise DocumentImportError(f"\u6587\u6863\u4e0d\u5b58\u5728\uff1a{file_path}")
        
        if not file_path.is_file():
            raise DocumentImportError(f"\u6307\u5b9a\u7684\u8def\u5f84\u4e0d\u662f\u4e00\u4e2a\u6587\u4ef6\uff1a{file_path}")

        # 检查文件大小，防止处理过大的文件
        try:
            file_size = file_path.stat().st_size
            max_size = getattr(self.config, "max_file_size", 50 * 1024 * 1024)  # 默认50MB
            if file_size > max_size:
                raise DocumentImportError(f"\u6587\u6863\u5927\u5c0f\u8d85\u8fc7\u6700\u5927\u9650\u5236\uff0c\u6700\u5927\u6587\u6863\u5927\u5c0f\u4e3a{max_size / (1024 * 1024):.1f}MB")
        except Exception as exc:
            # 文件大小检查失败，继续处理，但记录警告
            pass

        suffix = file_path.suffix.lower()
        if suffix in self.TEXT_SUFFIXES:
            content = self._load_text_file(file_path)
            self._update_last_load_state(file_path, content, "text_file")
            return content
        if suffix in self.DOCX_SUFFIXES:
            content = self._load_docx_file(file_path)
            self._update_last_load_state(file_path, content, "docx")
            return content
        if suffix in self.PDF_SUFFIXES:
            return self._load_pdf_file(file_path, pdf_options=pdf_options)

        supported = ", ".join(sorted(self.SUPPORTED_SUFFIXES))
        raise DocumentImportError(f"\u6682\u4e0d\u652f\u6301\u7684\u6587\u4ef6\u7c7b\u578b\uff1a{suffix or '\u65e0\u540e\u7f00'}\u3002\u652f\u6301\uff1a{supported}")

    def _load_text_file(self, path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return path.read_text(encoding="utf-8", errors="ignore")

    def _load_docx_file(self, path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise DocumentImportError("读取 DOCX 需要安装 python-docx。") from exc

        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        tables: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    tables.append(" | ".join(cells))

        content = [line for line in paragraphs + tables if line]
        if not content:
            raise DocumentImportError("DOCX 文件中未提取到可用文本。")
        return "\n".join(content)

    def _load_pdf_file(
        self,
        path: Path,
        pdf_options: PdfImportOptions | None = None,
    ) -> str:
        text_pages = self._extract_pdf_text_layer(path)
        if text_pages:
            content = self._clean_pdf_text_layer_document(text_pages)
            self._update_last_load_state(path, content, "pdf_text_layer")
            return content

        if not self.config.enable_pdf_ocr:
            raise DocumentImportError("PDF \u672a\u68c0\u6d4b\u5230\u7a33\u5b9a\u6587\u5b57\u5c42\uff0c\u4e14\u5f53\u524d\u914d\u7f6e\u672a\u542f\u7528 OCR\u3002")

        ocr_result = self._perform_pdf_ocr(path, pdf_options=pdf_options)
        self._update_last_load_state(
            path,
            ocr_result.text,
            "pdf_ocr",
            ocr_page_range=ocr_result.page_range_label,
            ocr_page_count=len(ocr_result.page_numbers),
            ocr_cache_hit=ocr_result.cache_hit,
        )
        return ocr_result.text

    def _extract_pdf_text_layer(self, path: Path) -> list[str]:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise DocumentImportError("读取 PDF 需要安装 pypdf。") from exc

        reader = PdfReader(str(path))
        total_pages = len(reader.pages)
        pages: list[str] = []
        for page_number, page in enumerate(reader.pages, start=1):
            extracted = page.extract_text() or ""
            extracted = self._clean_pdf_text_layer_page(
                extracted,
                page_number=page_number,
                total_pages=total_pages,
            )
            if extracted:
                pages.append(extracted)
        return pages

    def _clean_pdf_text_layer_page(
        self,
        text: str,
        page_number: int | None = None,
        total_pages: int | None = None,
    ) -> str:
        normalized = self._normalize_pdf_page_text(text)
        if not normalized:
            return ""

        cleaned_lines: list[str] = []
        for raw_line in normalized.splitlines():
            line = raw_line.strip()
            if not line or self._is_pdf_noise_line(line, page_number, total_pages):
                continue

            line = self._collapse_pdf_spacing(line)
            line = self._deduplicate_repeated_text(line)
            line = self._normalize_pdf_line(line)

            if not line or self._is_pdf_noise_line(line, page_number, total_pages):
                continue
            if cleaned_lines and cleaned_lines[-1] == line:
                continue

            cleaned_lines.append(line)

        merged_lines = self._merge_pdf_soft_wrapped_lines(cleaned_lines)
        return "\n".join(merged_lines).strip()

    @classmethod
    def _clean_pdf_text_layer_document(cls, pages: list[str]) -> str:
        merged_lines: list[str] = []
        for page in pages:
            for raw_line in page.splitlines():
                line = raw_line.strip()
                if not line:
                    continue
                line = cls._normalize_pdf_line(cls._deduplicate_repeated_text(line))
                if merged_lines and cls._should_merge_pdf_lines(merged_lines[-1], line):
                    combined = merged_lines[-1] + line
                    merged_lines[-1] = cls._normalize_pdf_line(cls._deduplicate_repeated_text(combined))
                else:
                    merged_lines.append(line)

        merged_lines = cls._format_pdf_cover_lines(merged_lines)
        return "\n".join(merged_lines).strip()

    @classmethod
    def _format_pdf_cover_lines(cls, lines: list[str]) -> list[str]:
        if not lines:
            return []

        normalized_lines = [cls._normalize_pdf_line(cls._deduplicate_repeated_text(line)) for line in lines if line.strip()]
        if not normalized_lines:
            return []

        title = normalized_lines[0]
        remainder: list[str] = normalized_lines[1:]
        formatted: list[str] = []

        if title.startswith("政府工作报告") and title != "政府工作报告":
            formatted.append("政府工作报告")
            trailing = title[len("政府工作报告"):].strip()
            if trailing:
                remainder.insert(0, trailing)
        else:
            formatted.append(title)

        if remainder and remainder[0].startswith(("——", "--")):
            split_lines = cls._split_pdf_cover_subtitle_author(remainder[0])
            formatted.extend(split_lines)
            remainder = remainder[1:]

        if len(formatted) >= 3:
            formatted[2] = cls._normalize_pdf_author_line(formatted[2])
        elif len(formatted) == 2:
            formatted[1] = cls._normalize_pdf_author_line(formatted[1]) if cls._looks_like_author_line(formatted[1]) else formatted[1]

        return formatted + remainder

    @classmethod
    def _split_pdf_cover_subtitle_author(cls, line: str) -> list[str]:
        line = cls._normalize_pdf_line(line)
        if not line:
            return []

        meeting_match = re.match(r"^(——.*?(?:会议上|大会上|开幕会上|闭幕会上))(.*)$", line)
        if meeting_match:
            subtitle = cls._normalize_pdf_line(meeting_match.group(1))
            author = cls._normalize_pdf_author_line(meeting_match.group(2))
            return [part for part in (subtitle, author) if part]

        for marker in cls._AUTHOR_ROLE_MARKERS:
            index = line.find(marker, 1)
            if index > 0:
                subtitle = cls._normalize_pdf_line(line[:index])
                author = cls._normalize_pdf_author_line(line[index:])
                return [part for part in (subtitle, author) if part]

        return [line]

    @classmethod
    def _looks_like_author_line(cls, line: str) -> bool:
        return any(marker in line for marker in cls._AUTHOR_ROLE_MARKERS)

    @classmethod
    def _normalize_pdf_author_line(cls, line: str) -> str:
        line = cls._normalize_pdf_line(line)
        if not line:
            return ""

        for marker in cls._AUTHOR_ROLE_MARKERS:
            pattern = rf"^({re.escape(marker)})\s*([A-Za-z\u4e00-\u9fff\u00b7]{{2,8}})$"
            match = re.match(pattern, line)
            if match:
                return f"{match.group(1)} {match.group(2)}"
        return line

    @staticmethod
    def _normalize_pdf_page_text(text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = text.replace("\u3000", " ").replace("\xa0", " ")
        text = re.sub(r"[\u2000-\u200b\u202f\u205f\u2060]+", " ", text)
        text = re.sub(r"[ \t\f\v]+", " ", text)
        return text.strip()

    @classmethod
    def _is_pdf_noise_line(
        cls,
        line: str,
        page_number: int | None = None,
        total_pages: int | None = None,
    ) -> bool:
        if any(pattern.match(line) for pattern in cls._PDF_NOISE_PATTERNS):
            return True

        lowered = line.lower()
        if "gov.cn" in lowered:
            return True
        if ("政府网" in line or "政府⽹" in line) and re.search(r"\d+\s*/\s*\d+$", line):
            return True
        if page_number is not None and total_pages is not None and line == f"{page_number}/{total_pages}":
            return True
        return False

    @classmethod
    def _collapse_pdf_spacing(cls, line: str) -> str:
        line = re.sub(r"\s+", " ", line).strip()
        line = re.sub(r"(?<=\d)\s+(?=\d)", "", line)
        line = re.sub(r"\s*([，。！？；：、（）《》“”‘’【】])\s*", r"\1", line)
        line = re.sub(r"\s*([—–-]{1,2})\s*", r"\1", line)
        line = re.sub(
            rf"(?<=[{cls._CJK_TOKEN}{cls._NUMERAL_TOKEN}])\s+(?=[{cls._CJK_TOKEN}{cls._NUMERAL_TOKEN}])",
            "",
            line,
        )
        line = re.sub(r"\s{2,}", " ", line)
        return line.strip()

    @classmethod
    def _deduplicate_repeated_text(cls, line: str) -> str:
        previous = None
        current = line
        while current != previous:
            previous = current
            current = cls._collapse_exact_repetition(current)
            current = cls._collapse_repeated_prefix(current)
            current = cls._collapse_repeated_phrases(current)
        return current

    @staticmethod
    def _collapse_exact_repetition(text: str) -> str:
        max_unit = min(len(text) // 2, 80)
        for unit_length in range(2, max_unit + 1):
            if len(text) % unit_length != 0:
                continue
            repeat_count = len(text) // unit_length
            if repeat_count < 2:
                continue
            unit = text[:unit_length]
            if unit * repeat_count == text:
                return unit
        return text

    @staticmethod
    def _collapse_repeated_prefix(text: str) -> str:
        max_prefix = min(len(text) // 2, 80)
        for prefix_length in range(max_prefix, 1, -1):
            prefix = text[:prefix_length]
            if text.startswith(prefix * 2):
                return prefix + text[prefix_length * 2 :]
        return text

    @classmethod
    def _collapse_repeated_phrases(cls, text: str) -> str:
        pattern = re.compile(r"([一-鿿0-9〇零一二三四五六七八九十百千万亿两]{2,40})\\1")
        collapsed = text
        while True:
            updated = pattern.sub(r"\\1", collapsed)
            if updated == collapsed:
                return updated
            collapsed = updated

    @classmethod
    def _normalize_pdf_line(cls, line: str) -> str:
        line = line.strip(" \t_")
        line = re.sub(r"\s*([，。！？；：、（）《》“”‘’【】])\s*", r"\1", line)
        line = re.sub(r"\s*([—–-]{1,2})\s*", r"\1", line)
        return line.strip()

    @classmethod
    def _merge_pdf_soft_wrapped_lines(cls, lines: list[str]) -> list[str]:
        merged: list[str] = []
        for line in lines:
            if merged and cls._should_merge_pdf_lines(merged[-1], line):
                merged[-1] = merged[-1] + line
            else:
                merged.append(line)
        return merged

    @classmethod
    def _should_merge_pdf_lines(cls, previous: str, current: str) -> bool:
        if not previous or not current:
            return False
        if previous.endswith(("。", "！", "？", "；", ":", "：")):
            return False
        if current.startswith(("——", "--")):
            return False
        if any(pattern.match(current) for pattern in cls._PDF_HEADING_PATTERNS):
            return False
        if len(previous) <= 20 and any(pattern.match(previous) for pattern in cls._PDF_HEADING_PATTERNS):
            return False
        if previous.endswith(("，", "、", "（", "(", "《", "“", "‘", "—", "-")):
            return True
        if re.search(rf"[{cls._CJK_TOKEN}{cls._NUMERAL_TOKEN}]$", previous) and re.match(
            rf"^[{cls._CJK_TOKEN}{cls._NUMERAL_TOKEN}]",
            current,
        ):
            return True
        return False

    @staticmethod
    def _normalize_ocr_page_spec(page_spec: str) -> str:
        normalized = (page_spec or "").strip()
        normalized = normalized.replace(chr(0xFF0C), ",").replace(chr(0x3001), ",")
        normalized = normalized.replace(chr(0xFF0D), "-").replace(chr(0x2014), "-").replace(chr(0x2013), "-")
        normalized = re.sub(r"\s+", "", normalized)
        return normalized

    @classmethod
    def _format_page_numbers(cls, page_numbers: list[int]) -> str:
        if not page_numbers:
            return ""
        ordered = sorted(dict.fromkeys(page_numbers))
        ranges: list[str] = []
        start = ordered[0]
        end = ordered[0]
        for value in ordered[1:]:
            if value == end + 1:
                end = value
                continue
            ranges.append(f"{start}-{end}" if start != end else str(start))
            start = value
            end = value
        ranges.append(f"{start}-{end}" if start != end else str(start))
        return ",".join(ranges)

    def _get_pdf_page_count(self, path: Path) -> int:
        try:
            import fitz
        except ImportError as exc:
            raise DocumentImportError("\u626b\u63cf\u7248 PDF OCR \u9700\u8981\u5b89\u88c5 PyMuPDF\u3002") from exc

        try:
            with fitz.open(str(path)) as document:
                return len(document)
        except Exception as exc:
            raise DocumentImportError(f"\u65e0\u6cd5\u8bfb\u53d6 PDF \u9875\u6570\uff1a{exc}") from exc

    def _resolve_ocr_page_numbers(
        self,
        path: Path,
        pdf_options: PdfImportOptions | None = None,
    ) -> list[int]:
        total_pages = self._get_pdf_page_count(path)
        if total_pages <= 0:
            raise DocumentImportError("PDF \u4e2d\u672a\u53d1\u73b0\u53ef\u7528\u9875\u9762\u3002")

        page_limit = max(1, int(getattr(self.config, "pdf_ocr_max_pages", 1) or 1))
        page_spec = self._normalize_ocr_page_spec(pdf_options.ocr_page_spec if pdf_options else "")
        if not page_spec:
            return list(range(1, min(total_pages, page_limit) + 1))

        selected_pages: list[int] = []
        seen: set[int] = set()
        for token in [item for item in page_spec.split(",") if item]:
            if "-" in token:
                start_text, end_text = token.split("-", 1)
                if not start_text.isdigit() or not end_text.isdigit():
                    raise DocumentImportError("OCR \u9875\u7801\u8303\u56f4\u683c\u5f0f\u65e0\u6548\uff0c\u8bf7\u4f7f\u7528 1-3,5 \u8fd9\u7c7b\u5199\u6cd5\u3002")
                start_page = int(start_text)
                end_page = int(end_text)
                if start_page > end_page:
                    raise DocumentImportError("OCR \u9875\u7801\u8303\u56f4\u65e0\u6548\uff0c\u8d77\u59cb\u9875\u4e0d\u80fd\u5927\u4e8e\u7ed3\u675f\u9875\u3002")
                page_values = range(start_page, end_page + 1)
            else:
                if not token.isdigit():
                    raise DocumentImportError("OCR \u9875\u7801\u8303\u56f4\u683c\u5f0f\u65e0\u6548\uff0c\u8bf7\u4f7f\u7528 1-3,5 \u8fd9\u7c7b\u5199\u6cd5\u3002")
                page_values = [int(token)]

            for page_number in page_values:
                if page_number < 1 or page_number > total_pages:
                    raise DocumentImportError(
                        f"OCR \u9875\u7801 {page_number} \u8d85\u51fa PDF \u5b9e\u9645\u9875\u6570 {total_pages}\u3002"
                    )
                if page_number not in seen:
                    seen.add(page_number)
                    selected_pages.append(page_number)

        if len(selected_pages) > page_limit:
            raise DocumentImportError(
                f"OCR \u9875\u7801\u8303\u56f4\u5171 {len(selected_pages)} \u9875\uff0c\u8d85\u8fc7\u5f53\u524d\u4e0a\u9650 pdf_ocr_max_pages={page_limit}\u3002"
            )
        return selected_pages

    def _should_use_ocr_cache(self, pdf_options: PdfImportOptions | None) -> bool:
        config_enabled = bool(getattr(self.config, "enable_ocr_result_cache", True))
        if not config_enabled:
            return False
        if pdf_options is None:
            return True
        return bool(pdf_options.use_ocr_cache)

    def _build_ocr_cache_key(self, path: Path, page_numbers: list[int]) -> str:
        stat = path.stat()
        payload = {
            "version": 1,
            "path": str(path.resolve()),
            "size": stat.st_size,
            "mtime_ns": stat.st_mtime_ns,
            "ocr_languages": self.config.ocr_languages,
            "pdf_ocr_zoom": float(getattr(self.config, "pdf_ocr_zoom", 1.0) or 1.0),
            "page_numbers": page_numbers,
        }
        encoded = json.dumps(payload, sort_keys=True, ensure_ascii=True).encode("utf-8")
        return hashlib.sha256(encoded).hexdigest()

    def _load_ocr_cache(self, cache_key: str) -> str | None:
        cache_dir = Path(self.config.resolved_ocr_cache_dir)
        cache_path = cache_dir / f"{cache_key}.json"
        if not cache_path.exists():
            return None
        try:
            payload = json.loads(cache_path.read_text(encoding="utf-8"))
        except (OSError, ValueError, json.JSONDecodeError):
            return None
        text = payload.get("text")
        return text if isinstance(text, str) and text.strip() else None

    def _write_ocr_cache(self, cache_key: str, result: OcrLoadResult) -> None:
        cache_dir = Path(self.config.resolved_ocr_cache_dir)
        cache_dir.mkdir(parents=True, exist_ok=True)
        cache_path = cache_dir / f"{cache_key}.json"
        payload = {
            "page_numbers": result.page_numbers,
            "page_range_label": result.page_range_label,
            "text": result.text,
        }
        cache_path.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def _perform_pdf_ocr(
        self,
        path: Path,
        pdf_options: PdfImportOptions | None = None,
    ) -> OcrLoadResult:
        page_numbers = self._resolve_ocr_page_numbers(path, pdf_options=pdf_options)
        page_range_label = self._format_page_numbers(page_numbers)
        use_cache = self._should_use_ocr_cache(pdf_options)
        cache_key = None
        if use_cache:
            cache_key = self._build_ocr_cache_key(path, page_numbers)
            cached_text = self._load_ocr_cache(cache_key)
            if cached_text:
                return OcrLoadResult(
                    text=cached_text,
                    page_numbers=page_numbers,
                    page_range_label=page_range_label,
                    cache_hit=True,
                )

        text = self._perform_pdf_ocr_uncached(path, page_numbers)
        result = OcrLoadResult(
            text=text,
            page_numbers=page_numbers,
            page_range_label=page_range_label,
            cache_hit=False,
        )
        if cache_key is not None:
            try:
                self._write_ocr_cache(cache_key, result)
            except OSError:
                pass
        return result

    def _perform_pdf_ocr_uncached(self, path: Path, page_numbers: list[int]) -> str:
        try:
            import fitz
        except ImportError as exc:
            raise DocumentImportError("\u626b\u63cf\u7248 PDF OCR \u9700\u8981\u5b89\u88c5 PyMuPDF\u3002") from exc

        try:
            from PIL import Image
        except ImportError as exc:
            raise DocumentImportError("\u626b\u63cf\u7248 PDF OCR \u9700\u8981\u5b89\u88c5 Pillow\u3002") from exc

        try:
            import pytesseract
            from pytesseract import TesseractNotFoundError
        except ImportError as exc:
            raise DocumentImportError(
                "\u626b\u63cf\u7248 PDF OCR \u9700\u8981\u5b89\u88c5 pytesseract\uff0c\u5e76\u5728\u7cfb\u7edf\u4e2d\u5b89\u88c5 Tesseract-OCR\u3002"
            ) from exc

        if self.config.tesseract_cmd:
            apply_tesseract_runtime_environment(self.config.tesseract_cmd)
            pytesseract.pytesseract.tesseract_cmd = self.config.tesseract_cmd

        pages: list[str] = []
        try:
            with fitz.open(str(path)) as document:
                zoom = max(float(self.config.pdf_ocr_zoom), 1.0)
                matrix = fitz.Matrix(zoom, zoom)

                for page_number in page_numbers:
                    page = document.load_page(page_number - 1)
                    pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                    image = Image.open(BytesIO(pixmap.tobytes("png")))
                    try:
                        text = pytesseract.image_to_string(
                            image,
                            lang=self.config.ocr_languages,
                        ).strip()
                    finally:
                        image.close()
                    if text:
                        pages.append(f"[OCR \u7b2c {page_number} \u9875]\n{text}")
        except TesseractNotFoundError as exc:
            raise DocumentImportError(
                "\u672a\u68c0\u6d4b\u5230 Tesseract-OCR \u53ef\u6267\u884c\u7a0b\u5e8f\u3002\u8bf7\u5b89\u88c5\u540e\uff0c\u6216\u5728\u914d\u7f6e\u4e2d\u8bbe\u7f6e tesseract_cmd\u3002"
            ) from exc
        except Exception as exc:
            raise DocumentImportError(f"OCR \u63d0\u53d6\u5931\u8d25\uff1a{exc}") from exc

        if not pages:
            raise DocumentImportError("PDF \u4e2d\u672a\u63d0\u53d6\u5230\u6587\u5b57\u5c42\uff0c\u4e14 OCR \u4e5f\u672a\u8bc6\u522b\u51fa\u53ef\u7528\u6587\u672c\u3002")
        return "\n\n".join(pages)
