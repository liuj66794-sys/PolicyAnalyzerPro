import shutil
import uuid
import unittest
from pathlib import Path

import fitz
from docx import Document as WordDocument

from core.algorithms import PolicyReportAnalyzer
from core.startup_checks import DeploymentCheck, StartupCheckReport
from importers.document_loader import DocumentLoader
from main import build_startup_check_payload


class RegressionSampleTests(unittest.TestCase):
    def setUp(self) -> None:
        base_root = Path("tests/_tmp_regression_samples")
        base_root.mkdir(parents=True, exist_ok=True)
        self.temp_root = base_root / f"regression_samples_{uuid.uuid4().hex}"
        self.temp_root.mkdir(parents=True, exist_ok=False)
        self.loader = DocumentLoader()
        self.analyzer = PolicyReportAnalyzer()

    def tearDown(self) -> None:
        shutil.rmtree(self.temp_root, ignore_errors=True)

    def _write_text_sample(self, path: Path) -> None:
        path.write_text(
            "政府工作报告\n\n"
            "十四届全国人大三次会议强调，"
            "要坚持稳中求进、以进促稳。\n\n"
            "2025年将继续扩大内需，加快发展新质生产力。\n",
            encoding="utf-8",
        )

    def _write_docx_sample(self, path: Path) -> None:
        document = WordDocument()
        document.add_heading("政策协同专题材料", level=1)
        document.add_paragraph("聚焦扩大内需、绿色转型与数字政务。")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "重点任务"
        table.cell(0, 1).text = "扩大内需"
        table.cell(1, 0).text = "工作进度"
        table.cell(1, 1).text = "按季度评估"
        document.save(path)

    def _write_text_layer_pdf(self, path: Path) -> None:
        font_path = Path("assets/fonts/simhei.ttf").resolve()
        document = fitz.open()
        page = document.new_page(width=595, height=842)
        page.insert_font(fontname="F0", fontfile=str(font_path))
        lines = [
            "政府工作报告",
            "十四届全国人大三次会议",
            "2025年重点任务：扩大内需，稳定预期，推进科技创新。",
            "要坚持高质量发展，推动产业升级和区域协同发展。",
        ]
        y = 90
        for line in lines:
            page.insert_text((72, y), line, fontsize=16, fontname="F0")
            y += 34
        document.save(path)
        document.close()

    def test_realistic_text_docx_and_pdf_samples_can_be_loaded(self) -> None:
        txt_path = self.temp_root / "sample_policy.txt"
        docx_path = self.temp_root / "sample_policy.docx"
        pdf_path = self.temp_root / "sample_policy.pdf"
        self._write_text_sample(txt_path)
        self._write_docx_sample(docx_path)
        self._write_text_layer_pdf(pdf_path)

        txt_content = self.loader.load_text_from_path(txt_path)
        docx_content = self.loader.load_text_from_path(docx_path)
        pdf_content = self.loader.load_text_from_path(pdf_path)

        self.assertIn("扩大内需", txt_content)
        self.assertIn("政策协同专题材料", docx_content)
        self.assertIn("重点任务 | 扩大内需", docx_content)
        self.assertIn("政府工作报告", pdf_content)
        self.assertIn("稳定预期", pdf_content)
        self.assertEqual(self.loader.last_load_state.extraction_mode, "pdf_text_layer")

    def test_realistic_pdf_sample_supports_metadata_and_structure_analysis(self) -> None:
        pdf_path = self.temp_root / "analysis_sample.pdf"
        self._write_text_layer_pdf(pdf_path)

        pdf_content = self.loader.load_text_from_path(pdf_path)
        prepared = self.analyzer.prepare_text(pdf_content)
        structure = self.analyzer.analyze_text_structure(prepared)

        self.assertIn("十四届全国人大三次会议", prepared.metadata["meeting_labels"])
        self.assertIn("2025年", prepared.metadata["years"])
        self.assertGreaterEqual(len(structure["paragraph_lengths"]), 3)
        self.assertGreater(structure["avg_paragraph_length"], 0)
        self.assertGreater(structure["longest_sentence_length"], 0)

    def test_startup_check_payload_serializes_for_cli_artifacts(self) -> None:
        report = StartupCheckReport(
            results=[
                DeploymentCheck(
                    key="config",
                    title="配置文件",
                    status="ok",
                    summary="默认配置文件可用。",
                ),
                DeploymentCheck(
                    key="ocr",
                    title="OCR",
                    status="warning",
                    summary="OCR 语言包需要复核。",
                    required=False,
                ),
            ],
            checked_at="2026-03-31T23:00:00+08:00",
        )

        payload = build_startup_check_payload(report)

        self.assertEqual(payload["overall_status"], "warning")
        self.assertEqual(payload["warning_count"], 1)
        self.assertEqual(payload["error_count"], 0)
        self.assertEqual(payload["results"][1]["summary"], "OCR 语言包需要复核。")


if __name__ == "__main__":
    unittest.main()
